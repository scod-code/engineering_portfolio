#!/usr/bin/env python3
"""
Create a test thesis .docx file to demonstrate the editor.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_thesis():
    """Create a test thesis document with common issues."""
    doc = Document()
    
    # Title
    title = doc.add_heading('AdaptLearn: Hyperheuristic-Orchestrated Intelligent Tutoring System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    doc.add_paragraph('COMP40321 Research Methods Submission')
    doc.add_paragraph('MSc in Robotics and Intelligent Systems')
    doc.add_paragraph('Supervised by Dr Jordan Bird')
    doc.add_paragraph('')
    
    # Section 2.2 Objectives (with issues)
    doc.add_heading('2.2 Objectives', level=1)
    
    # Problematic paragraphs with em-dashes and AI phrases
    objectives = [
        "Objective 1: Develop a multimodal perception module—which integrates computer vision and natural language processing—to capture learner engagement and cognitive states.",
        "Objective 2: Design a hyperheuristic orchestration framework; this framework will dynamically select and combine pedagogical strategies based on real-time learner data; it will optimize learning pathways through reinforcement learning.",
        "Objective 3: Implement a continual reinforcement learning agent—the agent will adapt tutoring content and difficulty—ensuring personalized learning experiences over Weeks 1-8 of the study.",
        "Objective 4: Create an evaluation protocol; the protocol will assess system efficacy using n = 50 participants; it will measure learning gains, engagement, and satisfaction.",
        "Objective 5: It is important to note that this research aims to leverage state-of-the-art machine learning techniques while harnessing multimodal data streams to deliver cutting-edge educational technology."
    ]
    
    for obj in objectives:
        doc.add_paragraph(obj)
    
    # Section 4 Resources Required
    doc.add_heading('4. Resources Required', level=1)
    
    resources = [
        "Computational compute: Access to cloud GPU instances (e.g., AWS p3.2xlarge) for model training—estimated budget: £500.",
        "Data: Anonymized learner interaction logs—to allow for transfer-learning across educational domains—from institutional partners.",
        "Supervision: Fort weekly meetings with Dr Jordan Bird to discuss progress and methodological refinements.",
        "Timeline: The project will run over Weeks 1-8, with specific milestones at Weeks 2-4 and 6-8."
    ]
    
    for res in resources:
        doc.add_paragraph(res)
    
    # More AI-sounding text
    doc.add_heading('Methodology', level=1)
    doc.add_paragraph("Furthermore, this study will utilize a mixed-methods approach. Moreover, it should be noted that the aforementioned techniques will be applied systematically. Additionally, the research seeks to contribute novel insights to the field of intelligent tutoring systems.")
    
    # Save
    filename = "test_thesis_sample.docx"
    doc.save(filename)
    print(f"✅ Created test thesis: {filename}")
    print(f"📏 This file contains all the issues the editor can fix:")
    print(f"   • Em-dashes (—) in multiple places")
    print(f"   • Semicolon chains")
    print(f"   • Week ranges with hyphens (Weeks 1-8)")
    print(f"   • AI phrases (leveraging, harnessing, utilizing)")
    print(f"   • Typos (fort weekly → fortnightly)")
    print(f"\n💡 Use this to test the editor before using your real thesis.")

if __name__ == "__main__":
    create_test_thesis()