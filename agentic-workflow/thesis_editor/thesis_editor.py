#!/usr/bin/env python3
"""
MSc Thesis Editor for Word Documents (.docx)
Specifically for COMP40321 Research Methods submission
Works directly with Microsoft Word .docx files
"""

import os
import re
import sys
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Tuple
import json
import webbrowser

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


class ThesisEditor:
    """Editor for MSc thesis Word documents."""
    
    def __init__(self, docx_path: str):
        self.docx_path = Path(docx_path)
        if not self.docx_path.exists():
            raise FileNotFoundError(f"Document not found: {docx_path}")
        
        # Load document
        self.doc = Document(str(self.docx_path))
        self.original_paragraphs = [p.text for p in self.doc.paragraphs]
        
        # Create backup filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        self.backup_path = self.docx_path.parent / f"{self.docx_path.stem}_backup_{timestamp}{self.docx_path.suffix}"
        
        # Save backup
        self.doc.save(str(self.backup_path))
        
        # Changes log
        self.changes_log = []
        self.current_section = None
        
        print(f"📄 Loaded: {self.docx_path}")
        print(f"📏 Pages: ~{len(self.doc.paragraphs) // 50} (estimated)")
        print(f"💾 Backup: {self.backup_path}")
    
    def analyze_document(self) -> Dict:
        """Analyze the document for common issues."""
        print("\n🔍 Analyzing document...")
        
        all_text = "\n".join([p.text for p in self.doc.paragraphs])
        
        issues = {
            "em_dashes": self._count_pattern(all_text, r'—|--|–'),
            "semicolons": self._count_pattern(all_text, r';'),
            "week_ranges": self._find_week_ranges(all_text),
            "ai_phrases": self._find_ai_phrases(all_text),
            "sections": self._find_sections(),
            "paragraph_count": len(self.doc.paragraphs),
            "word_count": sum(len(p.text.split()) for p in self.doc.paragraphs)
        }
        
        # Print summary
        print(f"  • Paragraphs: {issues['paragraph_count']}")
        print(f"  • Words: {issues['word_count']:,}")
        print(f"  • Em-dashes: {issues['em_dashes']['count']}")
        print(f"  • Semicolons: {issues['semicolons']['count']}")
        print(f"  • Week ranges to fix: {len(issues['week_ranges'])}")
        print(f"  • AI phrases detected: {len(issues['ai_phrases'])}")
        print(f"  • Sections found: {len(issues['sections'])}")
        
        return issues
    
    def _count_pattern(self, text: str, pattern: str) -> Dict:
        """Count occurrences of a pattern."""
        matches = list(re.finditer(pattern, text))
        return {
            "count": len(matches),
            "examples": [text[max(0, m.start()-30):min(len(text), m.end()+30)] 
                        for m in matches[:3]] if matches else []
        }
    
    def _find_week_ranges(self, text: str) -> List[Dict]:
        """Find week ranges that need en-dashes."""
        # Patterns like: Weeks 1-8, Week 1-8, weeks 1-8
        patterns = [
            r'(Weeks?)\s+(\d+)\s*[-–]\s*(\d+)',
            r'(\d+)\s*[-–]\s*(\d+)\s+(weeks?)',
        ]
        
        week_ranges = []
        for pattern in patterns:
            matches = list(re.finditer(pattern, text, re.IGNORECASE))
            for match in matches:
                week_ranges.append({
                    "match": match.group(),
                    "start": match.start(),
                    "end": match.end(),
                    "context": text[max(0, match.start()-50):min(len(text), match.end()+50)]
                })
        
        return week_ranges
    
    def _find_ai_phrases(self, text: str) -> List[Dict]:
        """Find AI-generated sounding phrases."""
        ai_patterns = [
            r'leveraging\b',
            r'harnessing\b',
            r'utilizing\b',
            r'it is important to note that',
            r'it should be noted that',
            r'as previously mentioned',
            r'as discussed above',
            r'this (?:paper|study|research) (?:aims to|will|seeks to)',
            r'state-of-the-art',
            r'cutting-edge',
            r'novel approach',
            r'innovative method',
            r'furthermore',
            r'moreover',
            r'additionally',
            r'in conclusion',
            r'in summary',
            r'the aforementioned',
        ]
        
        ai_phrases = []
        for pattern in ai_patterns:
            matches = list(re.finditer(pattern, text, re.IGNORECASE))
            for match in matches:
                ai_phrases.append({
                    "phrase": match.group(),
                    "start": match.start(),
                    "end": match.end(),
                    "context": text[max(0, match.start()-50):min(len(text), match.end()+50)]
                })
        
        return ai_phrases
    
    def _find_sections(self) -> List[Dict]:
        """Find document sections."""
        sections = []
        for i, paragraph in enumerate(self.doc.paragraphs):
            text = paragraph.text.strip()
            
            # Check if this looks like a section header
            if (len(text) < 100 and  # Not too long
                (text.isupper() or  # ALL CAPS
                 re.match(r'^\d+\.\d+\s+', text) or  # 2.2 Title
                 re.match(r'^\d+\.\s+', text) or  # 2. Title
                 re.match(r'^[A-Z][A-Za-z\s]+:$', text) or  # Title:
                 paragraph.style.name.startswith('Heading'))):  # Word heading style
                
                sections.append({
                    "index": i,
                    "text": text,
                    "paragraph": paragraph
                })
        
        return sections
    
    def fix_week_ranges(self) -> int:
        """Fix week ranges (replace hyphens with en-dashes)."""
        print("\n🔧 Fixing week ranges (e.g., Weeks 1-8 → Weeks 1–8)...")
        
        changes_made = 0
        
        for i, paragraph in enumerate(self.doc.paragraphs):
            original_text = paragraph.text
            
            # Replace hyphens with en-dashes in week ranges
            # Pattern: Weeks? 1-8 or 1-8 weeks?
            fixed_text = re.sub(
                r'(Weeks?\s+)(\d+)\s*-\s*(\d+)',
                r'\1\2–\3',
                original_text,
                flags=re.IGNORECASE
            )
            
            fixed_text = re.sub(
                r'(\d+)\s*-\s*(\d+)\s+(weeks?)',
                r'\1–\2 \3',
                fixed_text,
                flags=re.IGNORECASE
            )
            
            if fixed_text != original_text:
                paragraph.text = fixed_text
                changes_made += 1
                
                # Log change
                self.changes_log.append({
                    "type": "week_range_fix",
                    "paragraph": i,
                    "original": original_text,
                    "replacement": fixed_text
                })
                
                print(f"  ✅ Fixed: {original_text[:50]}... → {fixed_text[:50]}...")
        
        print(f"  📊 Fixed {changes_made} week ranges")
        return changes_made
    
    def fix_em_dashes_paragraph_by_paragraph(self) -> int:
        """Fix em-dashes paragraph by paragraph with user approval."""
        print("\n🔧 Fixing em-dashes paragraph by paragraph...")
        print("   Use ':' for clause separation, ',' for parentheticals")
        print("   Press Enter to skip, 'q' to quit")
        
        changes_made = 0
        paragraphs_processed = 0
        
        for i, paragraph in enumerate(self.doc.paragraphs):
            text = paragraph.text
            
            # Check if this paragraph has em-dashes
            if '—' in text or '--' in text or '–' in text:
                paragraphs_processed += 1
                
                print(f"\n{'='*60}")
                print(f"📝 Paragraph {i+1}/{len(self.doc.paragraphs)}:")
                print(f"{'='*60}")
                print(text[:200] + "..." if len(text) > 200 else text)
                print(f"{'='*60}")
                
                # Show em-dash locations
                em_dash_positions = []
                for match in re.finditer(r'—|--|–', text):
                    start = max(0, match.start() - 20)
                    end = min(len(text), match.end() + 20)
                    context = text[start:end]
                    em_dash_positions.append({
                        "position": match.start(),
                        "context": context
                    })
                
                if em_dash_positions:
                    print(f"\n🔍 Found {len(em_dash_positions)} em-dash(es):")
                    for j, pos in enumerate(em_dash_positions):
                        print(f"  {j+1}. ...{pos['context']}...")
                    
                    # Ask for action
                    action = input("\nAction? (c=colon, m=comma, s=skip, q=quit): ").lower().strip()
                    
                    if action == 'q':
                        print("Quitting em-dash fix...")
                        break
                    elif action == 's':
                        print("Skipping...")
                        continue
                    elif action in ['c', 'm']:
                        replacement = ':' if action == 'c' else ','
                        
                        # Replace all em-dashes in this paragraph
                        fixed_text = re.sub(r'—|--|–', replacement, text)
                        paragraph.text = fixed_text
                        changes_made += 1
                        
                        # Log change
                        self.changes_log.append({
                            "type": "em_dash_fix",
                            "paragraph": i,
                            "replacement": replacement,
                            "original": text,
                            "replacement": fixed_text
                        })
                        
                        print(f"✅ Replaced with '{replacement}'")
                    else:
                        print("Skipping...")
        
        print(f"\n📊 Processed {paragraphs_processed} paragraphs with em-dashes")
        print(f"📊 Made {changes_made} replacements")
        return changes_made
    
    def scan_ai_phrases(self) -> List[Dict]:
        """Scan for AI phrases and prepare for Superhumanizer."""
        print("\n🤖 Scanning for AI-generated phrases...")
        
        ai_phrases = []
        
        for i, paragraph in enumerate(self.doc.paragraphs):
            text = paragraph.text
            
            # Check for AI phrases
            patterns = [
                (r'leveraging\b', 'leveraging'),
                (r'harnessing\b', 'harnessing'),
                (r'utilizing\b', 'utilizing'),
                (r'it is important to note that', 'it is important to note that'),
                (r'it should be noted that', 'it should be noted that'),
                (r'state-of-the-art', 'state-of-the-art'),
                (r'cutting-edge', 'cutting-edge'),
            ]
            
            for pattern, phrase in patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    ai_phrases.append({
                        "paragraph_index": i,
                        "paragraph_text": text[:100] + "..." if len(text) > 100 else text,
                        "phrase": phrase,
                        "full_text": text
                    })
                    break  # Only log once per paragraph
        
        if ai_phrases:
            print(f"🔍 Found {len(ai_phrases)} paragraphs with AI phrases:")
            for j, phrase_info in enumerate(ai_phrases[:5]):  # Show first 5
                print(f"  {j+1}. Paragraph {phrase_info['paragraph_index']+1}: {phrase_info['paragraph_text']}")
            
            if len(ai_phrases) > 5:
                print(f"  ... and {len(ai_phrases) - 5} more")
            
            # Ask if user wants to open Superhumanizer
            print("\n🌐 Open Superhumanizer.ai to humanize these paragraphs?")
            print("   (You can copy-paste paragraphs to humanize them)")
            open_browser = input("   Open browser? (y/n): ").lower().strip()
            
            if open_browser == 'y':
                webbrowser.open("https://superhumanizer.ai")
                print("✅ Opened Superhumanizer.ai in your browser")
                print("💡 Tip: Copy paragraphs from above and paste into Superhumanizer")
        
        return ai_phrases
    
    def save_edited_version(self) -> str:
        """Save the edited document."""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        edited_path = self.docx_path.parent / f"{self.docx_path.stem}_edited_{timestamp}{self.docx_path.suffix}"
        
        # Save document
        self.doc.save(str(edited_path))
        
        # Save changes log
        log_path = edited_path.parent / f"{edited_path.stem}_changes.json"
        with open(log_path, 'w', encoding='utf-8') as f:
            json.dump(self.changes_log, f, indent=2)
        
        print(f"\n💾 Saved edited document: {edited_path}")
        print(f"📋 Changes log: {log_path}")
        print(f"📊 Total changes made: {len(self.changes_log)}")
        
        # Show summary
        change_types = {}
        for change in self.changes_log:
            change_type = change.get("type", "unknown")
            change_types[change_type] = change_types.get(change_type, 0) + 1
        
        print("\n📊 Changes by type:")
        for change_type, count in change_types.items():
            print(f"  • {change_type}: {count}")
        
        return str(edited_path)


def find_docx_files() -> List[Path]:
    """Find .docx files in current directory."""
    current_dir = Path.cwd()
    docx_files = list(current_dir.glob("*.docx"))
    
    # Also check for common thesis filenames
    common_names = [
        "thesis.docx", "dissertation.docx", "research.docx",
        "paper.docx", "submission.docx", "COMP40321.docx",
        "AdaptLearn.docx", "MSc_Thesis.docx"
    ]
    
    # Check parent directory too
    parent_dir = current_dir.parent
    for name in common_names:
        parent_file = parent_dir / name
        if parent_file.exists():
            docx_files.append(parent_file)
    
    return docx_files


def main():
    """Main interactive editor."""
    print("=" * 60)
    print("📚 MSc Thesis Editor for Word Documents (.docx)")
    print("=" * 60)
    
    # Find .docx files
    docx_files = find_docx_files()
    
    if docx_files:
        print("\n📁 Found .docx files:")
        for i, file in enumerate(docx_files):
            print(f"  {i+1}. {file.name} ({file.parent})")
        
        file_choice = input("\nSelect file number, or enter full path: ").strip()
        
        try:
            # Try to parse as number
            file_index = int(file_choice) - 1
            if 0 <= file_index < len(docx_files):
                docx_path = docx_files[file_index]
            else:
                docx_path = Path(file_choice)
        except ValueError:
            # Use as path
            docx_path = Path(file_choice)
    else:
        print("\n📁 No .docx files found in current directory.")
        docx_path = Path(input("Enter full path to your .docx file: ").strip())
    
    # Create editor
    try:
        editor = ThesisEditor(str(docx_path))
    except FileNotFoundError as e:
        print(f"❌ {e}")
        return
    except Exception as e:
        print(f"❌ Error loading document: {e}")
        print("   Make sure it's a valid .docx file and not corrupted.")
        return
    
    # Analyze document
    issues = editor.analyze_document()
    
    # Main menu
    while True:
        print("\n" + "=" * 60)
        print("📝 THESIS EDITING MENU")
        print("=" * 60)
        print("1. Fix week ranges (Weeks 1-8 → Weeks 1–8) - Automatic")
        print("2. Fix em-dashes paragraph by paragraph")
        print("3. Scan for AI phrases (opens Superhumanizer.ai)")
        print("4. Save edited version and exit")
        print("5. Exit without saving")
        print("=" * 60)
        
        choice = input("\nSelect option (1-5): ").strip()
        
        if choice == "1":
            changes = editor.fix_week_ranges()
            if changes > 0:
                print(f"✅ Fixed {changes} week ranges")
            else:
                print("✅ No week ranges needed fixing")
        
        elif choice == "2":
            print("\n🔧 Em-dash Fixing Instructions:")
            print("   • Use 'c' to replace with colon (:) for clause separation")
            print("   • Use 'm' to replace with comma (,) for parentheticals")
            print("   • Use 's' to skip a paragraph")
            print("   • Use 'q' to quit and return to menu")
            print("\nExample: 'The system—which uses AI—will adapt.'")
            print("         Parenthetical → replace with comma: 'The system, which uses AI, will adapt.'")
            
            confirm = input("\nStart fixing em-dashes? (y/n): ").lower().strip()
            if confirm == 'y':
                editor.fix_em_dashes_paragraph_by_paragraph()
        
        elif choice == "3":
            ai_phrases = editor.scan_ai_phrases()
            if not ai_phrases:
                print("✅ No AI phrases detected!")
        
        elif choice == "4":
            edited_path = editor.save_edited_version()
            print(f"\n✅ Your original file is untouched: {docx_path}")
            print(f"✅ Your edited version is ready: {edited_path}")
            print("\n🎓 Good luck with your MSc submission!")
            break
        
        elif choice == "5":
            print("❌ Exiting without saving changes.")
            print(f"💾 Your original file is still safe: {docx_path}")
            break
        
        else:
            print("❌ Invalid choice. Please enter 1-5.")


if __name__ == "__main__":
    # Check for python-docx
    try:
        import docx
        main()
    except ImportError:
        print("❌ Required package not installed.")
        print("   Please run: pip install python-docx")
        print("   Or run the setup commands:")
        print("   python -m venv venv")
        print("   venv\\Scripts\\activate")
        print("   pip install python-docx")
        input("\nPress Enter to exit...")