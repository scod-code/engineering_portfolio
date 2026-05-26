#!/usr/bin/env python3
"""
Simple Word Document Automation Example
Automates common Word document tasks using Python
"""

from prefect import flow, task
from datetime import datetime
import os
from pathlib import Path


@task
def create_word_document(filename: str, content: str) -> str:
    """Create a new Word document with content."""
    print(f"📄 Creating document: {filename}")
    
    try:
        # Try using python-docx library
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('Automated Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('=' * 50)
        
        # Add main content
        doc.add_paragraph(content)
        
        # Add footer
        doc.add_paragraph('=' * 50)
        doc.add_paragraph('This document was automatically generated.')
        
        # Save document
        doc.save(filename)
        print(f"✅ Created: {filename}")
        return filename
        
    except ImportError:
        # Fallback: Create a simple text file if python-docx not installed
        print("⚠️  python-docx not installed, creating text file instead")
        text_filename = filename.replace('.docx', '.txt')
        with open(text_filename, 'w', encoding='utf-8') as f:
            f.write(f"Automated Report\n")
            f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"{'=' * 50}\n\n")
            f.write(content)
            f.write(f"\n\n{'=' * 50}\n")
            f.write("This document was automatically generated.")
        print(f"✅ Created: {text_filename}")
        return text_filename


@task
def update_word_document(filename: str, section: str, new_content: str) -> str:
    """Update a specific section in an existing document."""
    print(f"✏️  Updating section '{section}' in: {filename}")
    
    try:
        from docx import Document
        
        doc = Document(filename)
        
        # Find and replace text (simple approach)
        found = False
        for paragraph in doc.paragraphs:
            if section in paragraph.text:
                paragraph.text = new_content
                found = True
                print(f"✅ Updated section: {section}")
                break
        
        if not found:
            # Add new section if not found
            doc.add_heading(section, level=2)
            doc.add_paragraph(new_content)
            print(f"✅ Added new section: {section}")
        
        doc.save(filename)
        return filename
        
    except ImportError:
        print("⚠️  python-docx not installed")
        return filename


@task
def generate_report(data: dict, template: str = "weekly") -> str:
    """Generate a formatted report from data."""
    print(f"📊 Generating {template} report...")
    
    # Generate report content based on template
    if template == "weekly":
        content = f"""
Weekly Progress Report
======================

Date: {datetime.now().strftime('%Y-%m-%d')}

Summary:
- Tasks Completed: {data.get('tasks_completed', 0)}
- Tasks In Progress: {data.get('tasks_in_progress', 0)}
- Pending Tasks: {data.get('pending_tasks', 0)}

Highlights:
{data.get('highlights', 'No highlights this week.')}

Challenges:
{data.get('challenges', 'No challenges reported.')}

Next Week's Goals:
{data.get('next_goals', 'Continue current progress.')}
"""
    elif template == "status":
        content = f"""
Project Status Report
=====================

Project: {data.get('project_name', 'N/A')}
Status: {data.get('status', 'N/A')}
Completion: {data.get('completion', '0%')}

Key Metrics:
- Budget Used: {data.get('budget_used', 'N/A')}
- Timeline: {data.get('timeline', 'N/A')}
- Risks: {data.get('risks', 'None identified')}

Notes:
{data.get('notes', 'No additional notes.')}
"""
    else:
        content = f"""
Report
======

{data.get('content', 'No content provided.')}
"""
    
    # Create filename
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"reports/{template}_report_{timestamp}.docx"
    
    # Create document
    result = create_word_document.fn(filename, content)
    return result


@task
def merge_documents(sources: list, output: str) -> str:
    """Merge multiple Word documents into one."""
    print(f"📑 Merging {len(sources)} documents into: {output}")
    
    try:
        from docx import Document
        
        merged_doc = Document()
        
        for source in sources:
            if os.path.exists(source):
                source_doc = Document(source)
                
                # Add source title
                merged_doc.add_heading(os.path.basename(source), level=2)
                
                # Copy all paragraphs
                for paragraph in source_doc.paragraphs:
                    merged_doc.add_paragraph(paragraph.text)
                
                print(f"✅ Added: {source}")
            else:
                print(f"⚠️  File not found: {source}")
        
        merged_doc.save(output)
        print(f"✅ Created merged document: {output}")
        return output
        
    except ImportError:
        print("⚠️  python-docx not installed")
        return output


@task
def add_table_to_document(filename: str, headers: list, data: list) -> str:
    """Add a table to a Word document."""
    print(f"📊 Adding table to: {filename}")
    
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document(filename)
        
        # Add table
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        # Add data rows
        for row_data in data:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                row_cells[i].text = str(value)
        
        doc.save(filename)
        print(f"✅ Added table to: {filename}")
        return filename
        
    except ImportError:
        print("⚠️  python-docx not installed")
        return filename


@flow(name="Word Document Automation")
def word_automation_workflow(
    action: str = "create",
    filename: str = "output/automated_document.docx",
    content: str = "This is an automated document.",
    data: dict = None
):
    """
    Main workflow for Word document automation.
    
    Args:
        action: create, update, report, merge, table
        filename: Output filename
        content: Document content
        data: Data for report generation
    """
    print("=" * 60)
    print("🚀 Starting Word Document Automation")
    print("=" * 60)
    
    # Ensure output directory exists
    Path("output").mkdir(exist_ok=True, parents=True)
    Path("reports").mkdir(exist_ok=True, parents=True)
    
    if action == "create":
        result = create_word_document(filename, content)
        
    elif action == "update":
        section = data.get("section", "Updated Section") if data else "Updated Section"
        new_content = data.get("new_content", content) if data else content
        result = update_word_document(filename, section, new_content)
        
    elif action == "report":
        template = data.get("template", "weekly") if data else "weekly"
        report_data = data or {
            "tasks_completed": 5,
            "tasks_in_progress": 2,
            "pending_tasks": 3,
            "highlights": "Completed major milestone",
            "challenges": "Resource constraints",
            "next_goals": "Finish remaining tasks"
        }
        result = generate_report(report_data, template)
        
    elif action == "merge":
        sources = data.get("sources", []) if data else []
        result = merge_documents(sources, filename)
        
    elif action == "table":
        headers = data.get("headers", ["Column 1", "Column 2"]) if data else ["Column 1", "Column 2"]
        table_data = data.get("data", [["Value 1", "Value 2"]]) if data else [["Value 1", "Value 2"]]
        result = add_table_to_document(filename, headers, table_data)
        
    else:
        print(f"❌ Unknown action: {action}")
        result = None
    
    print("=" * 60)
    print("✅ Word Document Automation Complete!")
    print(f"📄 Output: {result}")
    print("=" * 60)
    
    return result


# Example usage
if __name__ == "__main__":
    # Example 1: Create a simple document
    print("\n1. Creating a simple document...")
    word_automation_workflow(
        action="create",
        filename="output/simple_report.docx",
        content="This is a simple automated report.\n\nIt was created using Python and Prefect."
    )
    
    # Example 2: Generate a weekly report
    print("\n2. Generating a weekly report...")
    word_automation_workflow(
        action="report",
        filename="output/weekly_status.docx",
        data={
            "template": "weekly",
            "tasks_completed": 8,
            "tasks_in_progress": 3,
            "pending_tasks": 5,
            "highlights": "Launched new feature, improved performance by 20%",
            "challenges": "Team capacity limited, some dependencies delayed",
            "next_goals": "Complete feature testing, prepare for release"
        }
    )
    
    # Example 3: Create a document with a table
    print("\n3. Creating a document with table...")
    word_automation_workflow(
        action="table",
        filename="output/data_table.docx",
        data={
            "headers": ["Name", "Status", "Priority", "Due Date"],
            "data": [
                ["Task 1", "Complete", "High", "2024-01-15"],
                ["Task 2", "In Progress", "Medium", "2024-01-20"],
                ["Task 3", "Pending", "Low", "2024-01-25"],
                ["Task 4", "Complete", "High", "2024-01-10"]
            ]
        }
    )
    
    print("\n✅ All examples completed! Check the 'output' folder for results.")